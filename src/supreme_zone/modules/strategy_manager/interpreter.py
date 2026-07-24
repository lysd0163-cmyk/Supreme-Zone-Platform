from __future__ import annotations

from dataclasses import asdict, dataclass, field
from pathlib import Path
import json
import re
from typing import Any

from pypdf import PdfReader

try:  # pragma: no cover - optional runtime dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None  # type: ignore[assignment]


DEFAULT_HIERARCHY: tuple[str, ...] = ("D1", "H4", "H1", "M15")
CORE_TERM_ALIASES: tuple[tuple[str, str], ...] = (
    ("LIQUIDITY INTENT", "Liquidity Intent"),
    ("EXTERNAL LIQUIDITY", "External Liquidity"),
    ("DEEP LIQUIDITY", "Deep Liquidity"),
    ("ENGINEERED LIQUIDITY", "Engineered Liquidity"),
    ("RESTING LIQUIDITY", "Resting Liquidity"),
    ("INDUCEMENT", "Inducement"),
    ("ORIGIN", "Origin"),
    ("DISPLACEMENT", "Displacement"),
    ("REPRICING", "Repricing"),
    ("ORDER BLOCK", "Order Block"),
    ("FAIR VALUE GAP", "Fair Value Gap"),
    ("BOS", "BOS"),
    ("CHOCH", "CHoCH"),
    ("ICT", "ICT"),
    ("SOF", "SOF"),
    ("POL", "POL"),
    ("WYCKOFF", "Wyckoff"),
    ("MOMENTUM", "Momentum"),
    ("ORDER FLOW", "Order Flow"),
    ("PREMIUM", "Premium"),
    ("DISCOUNT", "Discount"),
    ("FRESHNESS", "Freshness"),
    ("UNCONSUMED", "Unconsumed"),
    ("EXECUTION PATH", "Execution Path"),
    ("AMIR", "AMIR"),
    ("GEOMETRIC INTERSECTION", "Geometric Intersection"),
    ("MULTI-TIMEFRAME", "Multi-Timeframe Synchronization"),
)


@dataclass(slots=True, frozen=True)
class StrategySection:
    heading: str
    kind: str
    lines: tuple[str, ...] = ()


@dataclass(slots=True, frozen=True)
class StrategyInterpretation:
    title: str
    version: str
    versions_found: tuple[str, ...]
    source_format: str
    source_path: str
    summary: str
    sections: tuple[StrategySection, ...]
    core_terms: tuple[str, ...]
    protocols: tuple[str, ...]
    laws: tuple[str, ...]
    definitions: tuple[str, ...]
    timeframes: tuple[str, ...]
    official_outputs: tuple[str, ...]
    is_supreme_zone_engine: bool
    analysis: dict[str, Any] = field(default_factory=dict)
    rules: dict[str, Any] = field(default_factory=dict)
    document: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class StrategyInterpreter:
    def interpret(self, path: str | Path, source_text: str | None = None, structured_payload: dict[str, Any] | None = None) -> StrategyInterpretation:
        strategy_path = Path(path)
        raw_text = source_text if source_text is not None else self._read_source_text(strategy_path)
        normalized = self._normalize_text(raw_text)
        structured_payload = structured_payload or {}

        title = self._extract_title(normalized, structured_payload, strategy_path)
        versions_found = self._extract_versions(normalized)
        version = self._select_version(versions_found)
        timeframes = self._extract_timeframes(normalized)
        sections = self._extract_sections(strategy_path, raw_text)
        core_terms = self._extract_core_terms(normalized)
        protocols = self._filter_terms(sections, ("PROTOCOL",))
        laws = self._filter_terms(sections, ("LAW",))
        definitions = self._filter_terms(sections, ("DEFINITION",))
        outputs = self._extract_official_outputs(normalized)
        analysis = self._analysis_payload(timeframes, structured_payload)
        rules = self._rules_payload(core_terms, outputs, timeframes)
        document = {
            "title": title,
            "version": version,
            "versions_found": list(versions_found),
            "source_format": strategy_path.suffix.lower().lstrip(".") or "txt",
            "source_path": str(strategy_path),
            "summary": self._summary(title, version, timeframes, core_terms, outputs),
            "sections_count": len(sections),
            "sections": [asdict(section) for section in sections],
            "core_terms": list(core_terms),
            "protocols": list(protocols),
            "laws": list(laws),
            "definitions": list(definitions),
            "timeframes": list(timeframes),
            "official_outputs": list(outputs),
            "analysis": analysis,
            "rules": rules,
            "is_supreme_zone_engine": bool(self._is_supreme_zone_engine(normalized, title, outputs, timeframes)),
        }
        return StrategyInterpretation(
            title=title,
            version=version,
            versions_found=versions_found,
            source_format=strategy_path.suffix.lower().lstrip(".") or "txt",
            source_path=str(strategy_path),
            summary=document["summary"],
            sections=sections,
            core_terms=core_terms,
            protocols=protocols,
            laws=laws,
            definitions=definitions,
            timeframes=timeframes,
            official_outputs=outputs,
            is_supreme_zone_engine=document["is_supreme_zone_engine"],
            analysis=analysis,
            rules=rules,
            document=document,
        )

    def _read_source_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".docx":
            return self._read_docx(path)
        return path.read_text(encoding="utf-8", errors="ignore")

    def _read_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    def _read_docx(self, path: Path) -> str:
        if Document is None:  # pragma: no cover - dependency guard
            raise RuntimeError("python-docx is required to read .docx strategy files")
        doc = Document(str(path))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = self._normalize_line(paragraph.text)
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = self._normalize_line(cell.text)
                    if text:
                        lines.append(text)
        return "\n".join(lines)

    @staticmethod
    def _normalize_text(text: str) -> str:
        return "\n".join(line.strip() for line in text.replace("\r", "\n").splitlines())

    @staticmethod
    def _normalize_line(text: str) -> str:
        return re.sub(r"\s+", " ", text or "").strip()

    def _extract_title(self, text: str, payload: dict[str, Any], path: Path) -> str:
        top = payload.get("name") or payload.get("title")
        if isinstance(top, str) and top.strip():
            return top.strip()
        match = re.search(r"SUPREME\s+ZONE\s+ENGINE", text, flags=re.IGNORECASE)
        if match:
            return "SUPREME ZONE ENGINE"
        return path.stem.replace("_", " ").strip() or "SUPREME ZONE ENGINE"

    def _extract_versions(self, text: str) -> tuple[str, ...]:
        versions = re.findall(r"v\d+(?:\.\d+)?", text, flags=re.IGNORECASE)
        normalized = []
        for version in versions:
            candidate = version.lower().replace("v", "v")
            if candidate not in normalized:
                normalized.append(candidate)
        return tuple(normalized)

    def _select_version(self, versions: tuple[str, ...]) -> str:
        if not versions:
            return "v1.0"
        def sort_key(item: str) -> tuple[int, int]:
            match = re.search(r"v(\d+)(?:\.(\d+))?", item)
            if not match:
                return (0, 0)
            major = int(match.group(1))
            minor = int(match.group(2) or 0)
            return (major, minor)
        return sorted(versions, key=sort_key)[-1]

    def _extract_timeframes(self, text: str) -> tuple[str, ...]:
        raw = []
        aliases = [
            (r"\bD1\b|\b1D\b|\b1\s*D\b", "D1"),
            (r"\bH4\b|\b4H\b|\b4\s*H\b", "H4"),
            (r"\bH1\b|\b1H\b|\b1\s*H\b", "H1"),
            (r"\bM15\b|\b15M\b|\b15\s*M\b", "M15"),
        ]
        for pattern, label in aliases:
            if re.search(pattern, text, flags=re.IGNORECASE):
                raw.append(label)
        ordered: list[str] = []
        for item in DEFAULT_HIERARCHY:
            if item in raw and item not in ordered:
                ordered.append(item)
        for item in raw:
            if item not in ordered:
                ordered.append(item)
        return tuple(ordered or DEFAULT_HIERARCHY)

    def _extract_sections(self, path: Path, raw_text: str) -> tuple[StrategySection, ...]:
        blocks = self._paragraph_blocks(path, raw_text)
        sections: list[StrategySection] = []
        current_heading = self._title_from_blocks(blocks) or "DOCUMENT"
        current_kind = self._classify_heading(current_heading)
        current_lines: list[str] = []

        for block in blocks:
            text = block["text"]
            if block["is_heading"]:
                if current_lines:
                    sections.append(StrategySection(heading=current_heading, kind=current_kind, lines=tuple(current_lines)))
                current_heading = text
                current_kind = self._classify_heading(text)
                current_lines = []
                continue
            current_lines.append(text)

        if current_lines:
            sections.append(StrategySection(heading=current_heading, kind=current_kind, lines=tuple(current_lines)))
        return tuple(sections)

    def _paragraph_blocks(self, path: Path, raw_text: str) -> list[dict[str, Any]]:
        suffix = path.suffix.lower()
        if suffix == ".docx" and Document is not None:
            doc = Document(str(path))
            blocks: list[dict[str, Any]] = []
            for paragraph in doc.paragraphs:
                text = self._normalize_line(paragraph.text)
                if not text:
                    continue
                style_name = paragraph.style.name if paragraph.style else ""
                is_heading = style_name.startswith("Heading") or self._looks_like_heading(text)
                blocks.append({"text": text, "is_heading": is_heading})
            return blocks
        blocks = []
        for line in self._normalize_text(raw_text).splitlines():
            text = self._normalize_line(line)
            if not text:
                continue
            blocks.append({"text": text, "is_heading": self._looks_like_heading(text)})
        return blocks

    @staticmethod
    def _title_from_blocks(blocks: list[dict[str, Any]]) -> str | None:
        for block in blocks[:8]:
            text = str(block.get("text", "")).strip()
            if re.search(r"SUPREME\s+ZONE\s+ENGINE", text, flags=re.IGNORECASE):
                return "SUPREME ZONE ENGINE"
        return None

    @staticmethod
    def _looks_like_heading(text: str) -> bool:
        stripped = text.strip()
        if not stripped or stripped.startswith("-"):
            return False
        if re.fullmatch(r"AND|OR|OF|TO|THE|A|AN", stripped, flags=re.IGNORECASE):
            return False
        if re.fullmatch(r"PART\s+\d+\s*/\s*\d+", stripped, flags=re.IGNORECASE):
            return True
        if re.search(r"\b(PROTOCOL|ENGINE|LAW|DEFINITION|VALIDATION|FORMALIZATION|MATHEMATICS|PROTOCOLS)\b", stripped, flags=re.IGNORECASE):
            return True
        letters = [char for char in stripped if char.isalpha()]
        if not letters:
            return False
        upper_ratio = sum(char.isupper() for char in letters) / len(letters)
        return upper_ratio >= 0.7 and len(stripped.split()) <= 12

    @staticmethod
    def _classify_heading(heading: str) -> str:
        upper = heading.upper()
        if "DEFINITION" in upper:
            return "definition"
        if "PROTOCOL" in upper:
            return "protocol"
        if "LAW" in upper:
            return "law"
        if upper.startswith("PART"):
            return "part"
        return "section"

    def _extract_core_terms(self, text: str) -> tuple[str, ...]:
        found = []
        normalized = text.upper()
        for alias, canonical in CORE_TERM_ALIASES:
            if alias in normalized and canonical not in found:
                found.append(canonical)
        return tuple(found)

    def _filter_terms(self, sections: tuple[StrategySection, ...], markers: tuple[str, ...]) -> tuple[str, ...]:
        results: list[str] = []
        for section in sections:
            heading = section.heading.upper()
            if any(marker in heading for marker in markers):
                if section.heading not in results:
                    results.append(section.heading)
        return tuple(results)

    def _extract_official_outputs(self, text: str) -> tuple[str, ...]:
        outputs: list[str] = []
        for label in ("OFFICIAL BUY ZONE", "OFFICIAL SELL ZONE"):
            if label in text.upper():
                outputs.append(label)
        return tuple(outputs)

    def _analysis_payload(self, timeframes: tuple[str, ...], payload: dict[str, Any]) -> dict[str, Any]:
        analysis = payload.get("analysis", {}) if isinstance(payload.get("analysis", {}), dict) else {}
        rules = payload.get("rules", {}) if isinstance(payload.get("rules", {}), dict) else {}
        return {
            "timeframes": list(analysis.get("timeframes") or rules.get("timeframes") or timeframes or DEFAULT_HIERARCHY),
            "minimum_candles": int(analysis.get("minimum_candles", rules.get("minimum_candles", 500))),
            "lookback_bars": int(analysis.get("lookback_bars", rules.get("lookback_bars", 120))),
            "zone_width_ratio": float(analysis.get("zone_width_ratio", rules.get("zone_width_ratio", 0.35))),
            "prefer_recent": bool(analysis.get("prefer_recent", rules.get("prefer_recent", True))),
            "bias": str(analysis.get("bias", rules.get("bias", "neutral"))).lower(),
            "strict_mode": True,
            "source": "strategy-interpreter",
        }

    def _rules_payload(self, core_terms: tuple[str, ...], official_outputs: tuple[str, ...], timeframes: tuple[str, ...]) -> dict[str, Any]:
        return {
            "timeframe_hierarchy": list(timeframes or DEFAULT_HIERARCHY),
            "core_terms": list(core_terms),
            "official_outputs": list(official_outputs),
            "requires_single_buy_zone": True,
            "requires_single_sell_zone": True,
            "requires_binary_validation": True,
            "requires_geometric_intersection": True,
            "requires_mtf_alignment": True,
            "requires_freshness": True,
            "requires_unconsumed": True,
            "requires_deterministic_reproducibility": True,
        }

    def _summary(self, title: str, version: str, timeframes: tuple[str, ...], core_terms: tuple[str, ...], outputs: tuple[str, ...]) -> str:
        tokens = [title, version]
        if timeframes:
            tokens.append("MTF: " + ", ".join(timeframes))
        if core_terms:
            tokens.append(f"terms={len(core_terms)}")
        if outputs:
            tokens.append("outputs=" + ", ".join(outputs))
        return " | ".join(tokens)

    def _is_supreme_zone_engine(self, text: str, title: str, outputs: tuple[str, ...], timeframes: tuple[str, ...]) -> bool:
        required_terms = {"LIQUIDITY INTENT", "ORIGIN", "DISPLACEMENT", "REPRICING", "ORDER BLOCK", "FAIR VALUE GAP", "AMIR"}
        upper = text.upper()
        found_terms = {alias for alias, _ in CORE_TERM_ALIASES if alias in upper}
        return bool(
            re.search(r"SUPREME\s+ZONE\s+ENGINE", title, flags=re.IGNORECASE)
            and len(required_terms.intersection(found_terms)) >= 5
            and {"OFFICIAL BUY ZONE", "OFFICIAL SELL ZONE"}.issubset(set(outputs))
            and set(DEFAULT_HIERARCHY).intersection(set(timeframes))
        )
