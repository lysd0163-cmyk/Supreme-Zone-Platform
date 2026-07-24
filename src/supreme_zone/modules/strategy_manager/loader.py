from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import yaml
from pypdf import PdfReader

from ...core.exceptions import StrategyError
from .interpreter import StrategyInterpreter
from .models import StrategyDefinition
from .validator import StrategyValidator

try:  # pragma: no cover - optional runtime dependency
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - optional runtime dependency
    Document = None  # type: ignore[assignment]


class StrategyLoader:
    def __init__(self) -> None:
        self.validator = StrategyValidator()
        self.interpreter = StrategyInterpreter()

    def load(self, path: str | Path) -> StrategyDefinition:
        strategy_path = Path(path)
        self.validator.validate_path(strategy_path)

        suffix = strategy_path.suffix.lower()
        structured_payload: dict[str, Any] | None = None
        source_text: str

        if suffix in {".yaml", ".yml"}:
            structured_payload = yaml.safe_load(strategy_path.read_text(encoding="utf-8")) or {}
            if not isinstance(structured_payload, dict):
                raise StrategyError("Strategy YAML must contain a mapping at the root")
            source_text = yaml.safe_dump(structured_payload, allow_unicode=True, sort_keys=False)
        elif suffix == ".json":
            structured_payload = json.loads(strategy_path.read_text(encoding="utf-8"))
            if not isinstance(structured_payload, dict):
                raise StrategyError("Strategy JSON must contain an object at the root")
            source_text = json.dumps(structured_payload, ensure_ascii=False, indent=2)
        elif suffix == ".pdf":
            source_text = self._pdf_text(strategy_path)
        elif suffix == ".docx":
            source_text = self._docx_text(strategy_path)
        else:
            source_text = strategy_path.read_text(encoding="utf-8", errors="ignore")

        interpretation = self.interpreter.interpret(strategy_path, source_text=source_text, structured_payload=structured_payload)

        if structured_payload is None:
            merged_payload: dict[str, Any] = {
                "name": interpretation.title,
                "version": interpretation.version,
                "source": source_text,
                "active": True,
            }
        else:
            merged_payload = dict(structured_payload)
            merged_payload.setdefault("name", interpretation.title)
            merged_payload.setdefault("version", interpretation.version)
            merged_payload.setdefault("source", source_text)
            merged_payload.setdefault("active", True)

        merged_analysis = dict(interpretation.analysis)
        structured_analysis = merged_payload.get("analysis") if isinstance(merged_payload.get("analysis"), dict) else {}
        merged_analysis.update(structured_analysis)
        merged_rules = dict(interpretation.rules)
        structured_rules = merged_payload.get("rules") if isinstance(merged_payload.get("rules"), dict) else {}
        merged_rules.update(structured_rules)

        merged_payload["analysis"] = merged_analysis
        merged_payload["rules"] = merged_rules
        merged_payload["interpreter"] = interpretation.to_dict()
        merged_payload["document"] = interpretation.document
        merged_payload["summary"] = interpretation.summary
        merged_payload["source_format"] = interpretation.source_format
        merged_payload["official_outputs"] = list(interpretation.official_outputs)
        merged_payload["core_terms"] = list(interpretation.core_terms)
        merged_payload["timeframes"] = list(interpretation.timeframes)
        merged_payload["is_supreme_zone_engine"] = interpretation.is_supreme_zone_engine

        self.validator.validate_interpreted_payload(merged_payload)

        name = str(merged_payload.get("name", interpretation.title))
        version = str(merged_payload.get("version", interpretation.version))
        active = bool(merged_payload.get("active", True))

        return StrategyDefinition(
            name=name,
            version=version,
            source_path=strategy_path,
            raw=merged_payload,
            interpretation=interpretation.to_dict(),
            active=active,
        )

    def _pdf_text(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    def _docx_text(self, path: Path) -> str:
        if Document is None:  # pragma: no cover - dependency guard
            raise StrategyError("python-docx is required to read .docx strategy files")
        doc = Document(str(path))
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        lines.append(text)
        return "\n".join(lines)
